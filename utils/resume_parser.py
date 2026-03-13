"""
Resume Parser Module
Extracts text from PDF and DOCX resume files
Beginner-friendly Python code for AI Resume Analyzer
"""

import PyPDF2
from docx import Document


def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF resume file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF in lowercase
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a PDF
        Exception: If there's an error reading the PDF
    """
    # Check if file exists
    try:
        # Check if file ends with .pdf
        if not file_path.lower().endswith('.pdf'):
            raise ValueError("File must be a PDF file (.pdf)")
        
        # Open and read PDF
        with open(file_path, 'rb') as pdf_file:
            # Create a PDF reader object
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Initialize empty string for all text
            all_text = ""
            
            # Loop through all pages in the PDF
            for page in pdf_reader.pages:
                # Extract text from each page
                page_text = page.extract_text()
                all_text += page_text
            
            # Convert to lowercase and return
            return all_text.lower()
    
    except FileNotFoundError:
        raise FileNotFoundError(f"File not found: {file_path}")
    except ValueError as e:
        raise ValueError(f"Invalid file format: {e}")
    except Exception as e:
        raise Exception(f"Error reading PDF file: {e}")


def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX resume file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX in lowercase
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a DOCX file
        Exception: If there's an error reading the DOCX file
    """
    # Check if file exists
    try:
        # Check if file ends with .docx
        if not file_path.lower().endswith(('.docx', '.doc')):
            raise ValueError("File must be a DOCX or DOC file")
        
        # Open the DOCX document
        doc = Document(file_path)
        
        # Initialize empty string for all text
        all_text = ""
        
        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"
        
        # Extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
                all_text += "\n"
        
        # Convert to lowercase and return
        return all_text.lower()
    
    except FileNotFoundError:
        raise FileNotFoundError(f"File not found: {file_path}")
    except ValueError as e:
        raise ValueError(f"Invalid file format: {e}")
    except Exception as e:
        raise Exception(f"Error reading DOCX file: {e}")


def extract_text_from_resume(file_path):
    """
    Extract text from a resume file (PDF or DOCX).
    Automatically detects the file format.
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        str: Extracted text from the resume in lowercase
        
    Raises:
        ValueError: If the file format is not supported
        FileNotFoundError: If the file doesn't exist
        Exception: If there's an error reading the file
    """
    # Check file extension
    file_lower = file_path.lower()
    
    if file_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_lower.endswith(('.docx', '.doc')):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")


# Example usage (uncomment to test)
# if __name__ == "__main__":
#     try:
#         # Extract from PDF
#         pdf_text = extract_text_from_pdf("resume.pdf")
#         print("PDF text extracted successfully")
#         print(pdf_text[:200])  # Print first 200 characters
#     except Exception as e:
#         print(f"Error: {e}")
#     
#     try:
#         # Extract from DOCX
#         docx_text = extract_text_from_docx("resume.docx")
#         print("DOCX text extracted successfully")
#         print(docx_text[:200])  # Print first 200 characters
#     except Exception as e:
#         print(f"Error: {e}")
